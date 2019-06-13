from os.path import abspath, join
import os

import docx

from custom_invitations import custom_invitations


def test_style_run():
    doc = docx.Document()
    paragraph = doc.add_paragraph()
    run = paragraph.add_run('this is a run of sample text')

    custom_invitations.style_run(run=run, is_bold=True, is_italic=True, font_size=12)
    assert run.font.bold
    assert run.font.italic
    # assert run.font.size == 12

    custom_invitations.style_run(run=run, is_bold=False, is_italic=False, font_size=15)
    assert not run.font.bold
    assert not run.font.italic
    # assert run.font.size == 15


def test_generate_invitations():
    path = join(abspath('.'), 'custom_invitations', 'test_files')
    custom_invitations.generate_invitations(textfile=join(path, 'test.txt'),
                                            filename=join(path, 'test.docx'))

    doc = docx.Document(join(path, 'test.docx'))
    with open(join(path, 'test.txt'), 'r') as name_list:
        # verify text for each individual invitation (each page is 6 paragraphs)
        for x, line in enumerate(name_list, start=0):
            name = line.strip()

            assert doc.paragraphs[(6 * x) + 0].text.strip() == \
                   'It would be a pleasure to have the company of'
            assert doc.paragraphs[(6 * x) + 1].text == name
            assert doc.paragraphs[(6 * x) + 2].text == \
                   'at 11101 Memory lane on the evening of'
            assert doc.paragraphs[(6 * x) + 3].text == 'April 31st'
            assert doc.paragraphs[(6 * x) + 4].text == "at 24 O'Clock"

    os.remove(join(path, 'test.docx'))
